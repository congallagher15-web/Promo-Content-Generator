
import io
from pathlib import Path
import streamlit as st
from docx import Document
from deep_translator import GoogleTranslator

st.set_page_config(page_title="Promo Content Generator - MVP", page_icon="🎰", layout="centered")

st.title("🎰 Promo Content Generator - MVP")
st.caption("Demo: upload ANY promo brief → receive promo content as a .docx.")

uploaded = st.file_uploader(
    "📎 Upload any promotion brief to demonstrate the flow (.docx, .doc, .pdf, .txt, .md, .html)",
    type=["docx", "doc", "pdf", "txt", "md", "html", "htm"],
    accept_multiple_files=False,
)

asset_path = Path(__file__).parent / "assets" / "22116_PSC_PragmaticDraw_ICE_V2.docx"
docx_bytes = asset_path.read_bytes()

if uploaded is not None:
    st.success("✅ Brief received. Delivering brand-approved content (embedded).")
    st.download_button(
        label="⬇️ Download: Pragmatic Play $20,000 Prize Draw (Approved).docx",
        data=docx_bytes,
        file_name="Pragmatic_Play_20000_Prize_Draw_Approved.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

    st.markdown("---")
    st.subheader("🌍 Translate")
    lang = st.selectbox("Choose a language", ["Italian", "Spanish", "Portuguese"], index=0)
    if st.button("Translate"):
        # Map to language codes
        target_map = {"Italian":"it", "Spanish":"es", "Portuguese":"pt"}
        target = target_map.get(lang, "it")

        # Read the embedded docx and translate paragraph-by-paragraph
        src_doc = Document(io.BytesIO(docx_bytes))
        out_doc = Document()

        def copy_para(text, style_name=None):
            p = out_doc.add_paragraph()
            run = p.add_run(text)
            # try to apply heading styles if detected
            if style_name and style_name.startswith("Heading"):
                try:
                    p.style = style_name
                except Exception:
                    # best-effort: map to standard names
                    if "1" in style_name: p.style = "Heading 1"
                    elif "2" in style_name: p.style = "Heading 2"
                    elif "3" in style_name: p.style = "Heading 3"

        translator = GoogleTranslator(source="auto", target=target)

        for para in src_doc.paragraphs:
            raw = para.text or ""
            if not raw.strip():
                out_doc.add_paragraph("")  # preserve spacing
                continue
            try:
                translated = translator.translate(raw)
            except Exception:
                translated = raw  # fail-safe: keep original
            copy_para(translated, getattr(para.style, "name", None))

        bio = io.BytesIO()
        out_doc.save(bio)
        bio.seek(0)

        st.download_button(
            label=f"⬇️ Download translated Word (.docx) — {lang}",
            data=bio.read(),
            file_name=f"Pragmatic_Play_20000_Prize_Draw_{lang}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

else:
    st.info("Drag & drop a file above to enable the download button.")
