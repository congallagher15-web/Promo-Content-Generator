
import io
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt

def _render_prize_block(prizes: List[dict]) -> str:
    if not prizes:
        return "- $1,000 Cash – 1\n- $300 Casino Bonus – 10\n- $80 Casino Bonus – 50"
    lines = []
    for p in prizes:
        prize = (p.get("prize","") or "").replace("  ", " ").strip()
        qty = p.get("quantity", 1)
        dash = "–"
        lines.append(f"- {prize} {dash} {qty}")
    return "\n".join(lines)

def _render_weekly_windows(windows: List[str]) -> str:
    if not windows:
        return "- See offer page for weekly windows."
    return "\n".join(f"- {w}" for w in windows)

def _info_popup_terms(parsed: Dict[str, Any], title: str) -> str:
    mechanic = parsed.get("mechanic","")
    cap = parsed.get("entry_cap_per_week","100")
    featured = parsed.get("featured","Casino games")
    start = parsed.get("start_date","")
    end = parsed.get("end_date","")
    weeks = _render_weekly_windows(parsed.get("weekly_windows_text",[]))
    prizes = parsed.get("prizes",[])
    prize_block = _render_prize_block(prizes)

    lines = []
    lines.append("**Terms & Conditions**\n")
    lines.append("**What is being offered**")
    lines.append(f"Players who complete a Casino challenge during the offer period can earn entries to the {title} and be in with a chance to win cash or Casino Bonus prizes.")
    lines.append("\n**When is the offer being conducted**")
    if start or end:
        s = start or 'the start date'
        e = end or 'the end date'
        lines.append(f"This offer runs from {s} until {e}.")
    lines.append("Weekly periods:")
    lines.append(weeks)
    lines.append("\n**Who is eligible to take part and how can you qualify**")
    lines.append("Offer is available to real-money PokerStars Casino players in eligible markets.")
    lines.append("Players must opt in via the Challenges Window. Progress only counts after opt-in.")
    lines.append(f"To qualify, {mechanic}")
    lines.append(f"Players can complete the challenge up to **{cap} times per weekly period**.")
    lines.append("\n**Wagering requirements and limitations by type of game**")
    lines.append(f"Only play on **{featured}** will count towards progress. No minimum bet requirement applies unless stated.")
    lines.append("\n**Claiming and redeeming the offer**")
    lines.append("Tickets are credited automatically upon each challenge completion and entered into the relevant Prize Draw.")
    if prize_block:
        lines.append("Prizes will be awarded as follows:")
        lines.append(prize_block)
    lines.append("\n**What else do you need to know**")
    lines.append("Cash prizes carry no wagering requirements and cannot be assigned or transferred.")
    lines.append("Casino Bonuses are valid for 72 hours unless otherwise stated and may require acceptance in certain markets.")
    lines.append("Game availability varies by device and location.")
    lines.append("See here for general promotional Terms & Conditions.")
    return "\n".join(lines)

def _build_sections(parsed: Dict[str, Any]) -> Dict[str, Dict[str, str]]:
    title = parsed.get("title","Prize Draw")
    mechanic = parsed.get("mechanic","Complete the weekly Casino challenge to receive 1 Prize Draw ticket.")
    cap = parsed.get("entry_cap_per_week","100")
    featured = parsed.get("featured","Casino games")

    qmc = {
        "promotion_title": title,
        "promotion_description": f"Win prizes by completing weekly challenges. {mechanic} Up to {cap} tickets per week.",
        "info_banner_live": (
            f"Earn tickets to the {title}\n"
            f"Get 1 ticket each time you complete the weekly challenge on {featured}\n"
            f"Players could win cash or Casino Bonuses"
        ),
        "info_banner_completed": "Any completed Challenge will appear here. Click on the 'Live' tab to view any currently available Challenge.",
        "info_popup_terms": _info_popup_terms(parsed, title),
    }

    qtc = {
        "task_title": title,
        "task_description": f"{mechanic} You can earn up to {cap} tickets per week.",
        "task_info_message": "Your progress will be tracked automatically after opt-in.",
        "warning_message": f"Only play on {featured} will count towards your progress.",
        "action_button_pre": "Opt-in",
        "action_button_post": "Play",
    }

    qrc = {
        "prize_content": f"Entry to the {title}",
        "reward_credited_content": "Ticket credited",
    }

    return {
        "quest_multi_content": qmc,
        "quest_task_content": qtc,
        "quest_reward_content": qrc,
    }

def generate_markdown_preview(parsed: Dict[str, Any]) -> str:
    sections = _build_sections(parsed)
    qmc = sections["quest_multi_content"]
    qtc = sections["quest_task_content"]
    qrc = sections["quest_reward_content"]
    parts = []
    parts.append("# Quest Multi Content\n")
    parts.append("## Promotion Title\n")
    parts.append(qmc['promotion_title'] + "\n\n")
    parts.append("## Promotion Description\n")
    parts.append(qmc['promotion_description'] + "\n\n")
    parts.append("## Info Banner (Live tab)\n")
    parts.append(qmc['info_banner_live'] + "\n\n")
    parts.append("## Info Banner (Completed tab)\n")
    parts.append(qmc['info_banner_completed'] + "\n\n")
    parts.append("## Info Pop-up\n")
    parts.append(qmc['info_popup_terms'] + "\n\n")
    parts.append("# Quest Task Content\n\n")
    parts.append("## Task Title\n")
    parts.append(qtc['task_title'] + "\n\n")
    parts.append("## Task Description\n")
    parts.append(qtc['task_description'] + "\n\n")
    parts.append("## Task Info Message\n")
    parts.append(qtc['task_info_message'] + "\n\n")
    parts.append("## Warning Message\n")
    parts.append(qtc['warning_message'] + "\n\n")
    parts.append("## Action Button (pre opt-in)\n")
    parts.append(qtc['action_button_pre'] + "\n\n")
    parts.append("## Action Button (post opt-in)\n")
    parts.append(qtc['action_button_post'] + "\n\n")
    parts.append("# Quest Reward Content\n\n")
    parts.append("## Prize Content\n")
    parts.append(qrc['prize_content'] + "\n\n")
    parts.append("## Reward Credited Content\n")
    parts.append(qrc['reward_credited_content'])
    return "".join(parts)

def generate_content_docx(parsed: Dict[str, Any]) -> bytes:
    sections = _build_sections(parsed)
    d = Document()

    def add_heading(text, lvl=1):
        return d.add_heading(text, level=lvl)

    def add_para(text):
        p = d.add_paragraph(text)
        for run in p.runs:
            run.font.size = Pt(11)

    add_heading("Quest Multi Content", 1)
    qmc = sections["quest_multi_content"]
    add_heading("Promotion Title", 2); add_para(qmc["promotion_title"])
    add_heading("Promotion Description", 2); add_para(qmc["promotion_description"])
    add_heading("Info Banner (Live tab)", 2); add_para(qmc["info_banner_live"])
    add_heading("Info Banner (Completed tab)", 2); add_para(qmc["info_banner_completed"])
    add_heading("Info Pop-up", 2); add_para(qmc["info_popup_terms"])

    add_heading("Quest Task Content", 1)
    qtc = sections["quest_task_content"]
    add_heading("Task Title", 2); add_para(qtc["task_title"])
    add_heading("Task Description", 2); add_para(qtc["task_description"])
    add_heading("Task Info Message", 2); add_para(qtc["task_info_message"])
    add_heading("Warning Message", 2); add_para(qtc["warning_message"])
    add_heading("Action Button (pre opt-in)", 2); add_para(qtc["action_button_pre"])
    add_heading("Action Button (post opt-in)", 2); add_para(qtc["action_button_post"])

    add_heading("Quest Reward Content", 1)
    qrc = sections["quest_reward_content"]
    add_heading("Prize Content", 2); add_para(qrc["prize_content"])
    add_heading("Reward Credited Content", 2); add_para(qrc["reward_credited_content"])

    bio = io.BytesIO()
    d.save(bio)
    bio.seek(0)
    return bio.read()
